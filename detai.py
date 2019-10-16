import argparse
import pytesseract
from PIL import Image
from PyPDF2 import PdfFileReader
from imutils import contours
from pdf2image import convert_from_path
import skew
import DetectTable
import handleTable
import os
import imutils
import cv2
from docx import Document
from PdfToImages import pdfToImage

def writeToTxt(result, docName):
    if os.path.exists(docName) == False:
        dc = Document()
        dc.save(docName)
    # else:
    #     i =1
    #     while (True) :
    #         if os.path.exists(docName[:len(docName)-5]+str(i)+".docx")==False:
    #             docName = docName[:len(docName)-5]+str(i)+".docx"
    #             dc = Document()
    #             dc.save(docName)
    #             break
    #         i = i + 1
    document = Document(docName)
    for line in result:
        para = document.add_paragraph(line)
        paragraph_format = para.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
        paragraph_format.line_spacing = 1
    document.save(docName)

# param: bounding_boxs = [('',x,y,w,h,0), ('',x,y,w,h,1),...] 0 = bounding box of text, 1 = bounding box of table
def handleFile(fileName,deblur,handleTableBasic,handleTableAdvance):
    """
    :param fileName: name of image to be converted
    :param outputName: name of doc file to be saved
    :return:

    detect table and layout-analyzing
    """
    img = cv2.imread(fileName)
    # handle skew
    img = skew.skewImage(img)
    # handle table with not auto fill
    if handleTableBasic or handleTableAdvance:
        if handleTableBasic:
            mask = DetectTable.detectTable(img).run(1)
        else:
            mask = DetectTable.detectTable(img).run(2)
        maskName = "mask.jpg"
        mask_img = cv2.imread(maskName)
        ## resize
        (h, w,_) = mask_img.shape
        mask_img = imutils.resize(mask_img, width=w * 2, height=h * 2)
        listResult, listBigBox = handleTable.getTableCoordinate(mask_img)
        img = cv2.resize(img, (mask_img.shape[1], mask_img.shape[0]))
        # origin = img.copy()
        resultTable = handleTable.retreiveTextFromTable(listResult,img)
        for pt in listBigBox:
            (x, y, w, h) = pt
            img[y:(y + h - 1), x:(x + w - 1)] = 255
        # out, result = handleTable.process_par(img, origin, listBigBox) ## use for layout
    resultNotTable = pytesseract.image_to_string(img,lang="vie")
    return resultNotTable,resultTable

def saveResult(folder,resultNotTable,resultTable,choose=True):
    if choose:
        f = open(folder+"/text.txt","w+")
    else:
        f = open(folder+"/text.txt","a+")
    f.write(resultNotTable[:])
    f.close()
    f = open("./saved/"+uuid+"/text.txt","a+")
    k = 0
    for rs in resultTable:
        if k%4==0:
            f.write("\n")
        f.write(rs+" ")
        k = k +1
    f.close()

def preprocessFile(fileType,fileName,folder):
    names = []
    if fileType == "pdf":
        pdfToImage(fileName,folder)
        count = len(os.listdir(folder))
        for k in range(1,count+1):
            names.append(str(k)+".jpg")
    else:
        names = os.listdir(folder)
    i = 0
    for filename in names:
        filename = os.path.join(folder,filename)
        if ".jpg" in filename:
            resultNotTable,resultTable = handleFile(filename,1,1,0)
            if i==0:
                saveResult(folder,resultNotTable,resultTable)
            else:
                saveResult(folder,resultNotTable,resultTable,False)
        elif "docx" or "doc" or "txt" in fileName:
            f = open(fileName,"r")
            des = open(folder+"/text.txt","w+")
            des.write(f.read())
            des.close()
            f.close()
        i = i +1


if __name__ == '__main__':
    uuid = "1"
    folder = "./saved/"+uuid
    preprocessFile("pdf","./101.pdf","./saved/1/")