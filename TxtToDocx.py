from docx import Document
from sys import argv
import os.path
def txtToDocx(txt_file):
    f=open(txt_file)
    contents = f.read()
    f.close()
    if os.path.exists('output.docx') == False:
        dc = Document()
        dc.save('output.docx')
    document = Document('output.docx')
    for line in contents.splitlines():
        para = document.add_paragraph(line)
        paragraph_format = para.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
        paragraph_format.line_spacing = 1
    document.save('output.docx')